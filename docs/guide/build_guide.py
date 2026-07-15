from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "68737D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "CBD4DC"
WHITE = "FFFFFF"
INK = "20262C"


def set_cell_text(cell, text, bold=False, color=INK, size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    add_inline(paragraph, text, bold_default=bold, color=color, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text, bold_default=False, color=INK, size=11):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", size - 0.5, color="2B566E")
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EDF2F5")
            run._element.get_or_add_rPr().append(shading)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, bold=bold_default, color=color)


def set_paragraph_border(paragraph, side, color, size=12, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_field(paragraph, instruction, placeholder=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    return run


def add_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=-1) + 1
    next_num = max(existing_num, default=0) + 1

    def create(kind, abs_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
        level.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        level.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        level.append(p_pr)
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_id = OxmlElement("w:abstractNumId")
        abstract_id.set(qn("w:val"), str(abs_id))
        num.append(abstract_id)
        numbering.append(num)

    create("bullet", next_abs, next_num)
    create("decimal", next_abs + 1, next_num + 1)
    return next_num, next_num + 1


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def load_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def fit_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, title, subtitle, fill, border_color=NAVY):
    draw.rounded_rectangle(xy, radius=24, fill=f"#{fill}", outline=f"#{border_color}", width=3)
    title_font = load_font(30, True)
    body_font = load_font(22)
    x1, y1, x2, y2 = xy
    title_lines = fit_text(draw, title, title_font, x2 - x1 - 44)
    body_lines = fit_text(draw, subtitle, body_font, x2 - x1 - 44)
    y = y1 + 24
    for line in title_lines:
        draw.text((x1 + 22, y), line, font=title_font, fill=f"#{NAVY}")
        y += 36
    y += 8
    for line in body_lines:
        draw.text((x1 + 22, y), line, font=body_font, fill=f"#{INK}")
        y += 29


def arrow(draw, start, end):
    draw.line([start, end], fill=f"#{BLUE}", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 16, y - 9), (x - 16, y + 9)], fill=f"#{BLUE}")


def create_diagrams(output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    labels = {
        "users": "Kullan\u0131c\u0131lar",
        "clients": "Web taray\u0131c\u0131s\u0131 ve Tauri masa\u00fcst\u00fc istemcisi",
        "api": "DocSys API",
        "api_sub": "Kimlik, yetki, sat\u0131rlar, analiz ve ya\u015fam d\u00f6ng\u00fcs\u00fc",
        "collab": "Collaboration",
        "collab_sub": "Yjs, Hocuspocus ve zengin metin odalar\u0131",
        "worker": "Worker",
        "worker_sub": "Export, purge ve snapshot s\u0131k\u0131\u015ft\u0131rma",
        "data": "Veri servisleri",
        "data_sub": "PostgreSQL, Redis ve MinIO/S3",
        "login": "Giri\u015f",
        "login_sub": "Kullan\u0131c\u0131 ad\u0131/e-posta, parola, iste\u011fe ba\u011fl\u0131 sunucu",
        "config": "\u0130stemci yap\u0131land\u0131rmas\u0131",
        "config_sub": "API adresi ve collaboration URL ke\u015ffi",
        "auth": "Kimlik do\u011frulama",
        "auth_sub": "Web: HTTP-only cookie | Desktop: session Bearer token",
        "workspace": "\u00c7al\u0131\u015fma alan\u0131",
        "workspace_sub": "Organizasyon, dok\u00fcman a\u011fac\u0131 ve izinli i\u00e7erik",
    }

    image = Image.new("RGB", (1600, 900), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    draw_box(draw, (70, 335, 360, 555), labels["users"], labels["clients"], LIGHT_BLUE)
    draw_box(draw, (490, 95, 920, 285), labels["api"], labels["api_sub"], "EAF3F8")
    draw_box(draw, (490, 355, 920, 545), labels["collab"], labels["collab_sub"], "EEF5EA")
    draw_box(draw, (490, 615, 920, 805), labels["worker"], labels["worker_sub"], "FFF5E6")
    draw_box(draw, (1090, 335, 1530, 555), labels["data"], labels["data_sub"], LIGHT_GRAY)
    arrow(draw, (360, 420), (490, 190))
    arrow(draw, (360, 445), (490, 450))
    arrow(draw, (360, 470), (490, 710))
    arrow(draw, (920, 190), (1090, 400))
    arrow(draw, (920, 450), (1090, 445))
    arrow(draw, (920, 710), (1090, 500))
    system_path = output_dir / "system-context.png"
    image.save(system_path, dpi=(180, 180))

    image = Image.new("RGB", (1600, 900), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    boxes = [
        (60, 315, 345, 565, labels["login"], labels["login_sub"], LIGHT_BLUE),
        (455, 315, 740, 565, labels["config"], labels["config_sub"], "EAF3F8"),
        (850, 315, 1135, 565, labels["auth"], labels["auth_sub"], "EEF5EA"),
        (1245, 315, 1530, 565, labels["workspace"], labels["workspace_sub"], LIGHT_GRAY),
    ]
    for box in boxes:
        draw_box(draw, box[:4], box[4], box[5], box[6])
    arrow(draw, (345, 440), (455, 440))
    arrow(draw, (740, 440), (850, 440))
    arrow(draw, (1135, 440), (1245, 440))
    login_path = output_dir / "login-flow.png"
    image.save(login_path, dpi=(180, 180))
    return {"system_context": system_path, "login_flow": login_path}


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, "DocSys | U\u00e7tan Uca Mimari, \u0130\u015fletim ve Kullan\u0131m K\u0131lavuzu", color=MUTED, size=8.5)
    set_paragraph_border(p, "bottom", BORDER, 6, 3)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(p, "DocSys  |  ", color=MUTED, size=8.5)
    run = add_field(p, "PAGE", "1")
    set_font(run, size=8.5, color=MUTED)


def add_cover(document, toc_entries):
    for _ in range(4):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    add_inline(kicker, "\u00dcR\u00dcN VE \u0130\u015eLET\u0130M EL K\u0130TABI", bold_default=True, color=BLUE, size=11)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    add_inline(title, "DocSys", bold_default=True, color=NAVY, size=34)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    add_inline(subtitle, "U\u00e7tan Uca Mimari, \u0130\u015fletim ve Kullan\u0131m K\u0131lavuzu", color=DARK_BLUE, size=18)
    detail = document.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail.paragraph_format.space_after = Pt(72)
    add_inline(detail, "Web, Tauri masa\u00fcst\u00fc, API, collaboration, worker, veri, g\u00fcvenlik ve kullan\u0131c\u0131 ak\u0131\u015flar\u0131", color=MUTED, size=11)
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(4)
    add_inline(version, "S\u00fcr\u00fcm 1.0 | 15 Temmuz 2026", bold_default=True, color=NAVY, size=11)
    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(status, "D\u00fczenlenebilir kaynak: docs/guide/DOCSYS-UCTAN-UCA-KILAVUZ.md", color=MUTED, size=9.5)
    document.add_page_break()
    toc_heading = document.add_paragraph("\u0130\u00e7indekiler", style="Heading 1")
    toc_heading.paragraph_format.space_before = Pt(0)
    for index, (label, page) in enumerate(toc_entries):
        if index == 18:
            document.add_page_break()
            continuation = document.add_paragraph("\u0130\u00e7indekiler - devam", style="Heading 1")
            continuation.paragraph_format.space_before = Pt(0)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2.5)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_inline(paragraph, f"{label}\t{page}", color=INK, size=9.2)
    document.add_page_break()


def add_note(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(paragraph, LIGHT_GRAY)
    set_paragraph_border(paragraph, "left", BLUE, 18, 8)
    add_inline(paragraph, text, color=DARK_BLUE, size=10.3)


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_paragraph_shading(paragraph, "F1F4F6")
    set_paragraph_border(paragraph, "left", "8CA8B8", 10, 5)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line or " ")
        set_font(run, "Consolas", 8.2, color="253845")


def parse_table(lines):
    rows = []
    for line in lines:
        values = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(values)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", value.replace(" ", "")) for value in rows[1]):
        rows.pop(1)
    return rows


def column_widths(rows):
    column_count = max(len(row) for row in rows)
    weights = []
    for index in range(column_count):
        maximum = max((len(row[index]) if index < len(row) else 0) for row in rows)
        weights.append(max(8, min(45, maximum)))
    total = sum(weights)
    widths = [round(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document, rows):
    if not rows:
        return
    widths = column_widths(rows)
    table = document.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for column_index in range(len(widths)):
            value = values[column_index] if column_index < len(values) else ""
            cell = row.cells[column_index]
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
                set_cell_text(cell, value, bold=True, color=NAVY, size=9.3)
            else:
                set_cell_text(cell, value, color=INK, size=9.1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)


def add_diagram(document, path, alt_text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.35))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(9)
    add_inline(caption, alt_text, color=MUTED, size=8.5)


def build(source, output, asset_dir):
    lines = source.read_text(encoding="utf-8").splitlines()
    headings = [re.match(r"^# (\d+\. .+)$", line).group(1) for line in lines if re.match(r"^# (\d+\. .+)$", line)]
    page_numbers = [4, 6, 7, 9, 11, 13, 14, 15, 17, 18, 19, 19, 20, 20, 21, 22, 23, 24, 25, 25, 26, 27, 28, 29, 30, 31, 32, 33, 33, 35, 36, 37, 37, 39, 39]
    toc_entries = list(zip(headings, page_numbers, strict=True))
    document = Document()
    configure_styles(document)
    configure_section(document)
    bullet_num, decimal_num = add_numbering(document)
    add_cover(document, toc_entries)
    diagrams = create_diagrams(asset_dir)
    in_code = False
    code_lines = []
    index = 0
    first_source_title = True
    source_body_started = False
    while index < len(lines):
        line = lines[index]
        if not source_body_started:
            if line == "## Dok\u00fcman\u0131n amac\u0131 ve hedef kitlesi":
                source_body_started = True
            else:
                index += 1
                continue
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("Bu Markdown dosyas\u0131"):
            index += 1
            continue
        if line == "[PAGEBREAK]":
            document.add_page_break()
            index += 1
            continue
        diagram_match = re.fullmatch(r"\[DIAGRAM:([a-z_]+)\]", line)
        if diagram_match:
            key = diagram_match.group(1)
            labels = {
                "system_context": "Sekil: DocSys sistem baglami ve servis iliskileri",
                "login_flow": "Sekil: Web ve masaustu giris akisi",
            }
            add_diagram(document, diagrams[key], labels[key])
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, parse_table(table_lines))
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and first_source_title:
                first_source_title = False
            else:
                paragraph = document.add_paragraph(style=f"Heading {level}")
                add_inline(paragraph, text, bold_default=True, color=BLUE if level < 3 else DARK_BLUE, size={1: 16, 2: 13, 3: 12}[level])
            index += 1
            continue
        if line.startswith("> "):
            add_note(document, line[2:])
            index += 1
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num)
            add_inline(paragraph, line[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, decimal_num)
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        if re.match("^(S\u00fcr\u00fcm|Tarih|Kapsam):", line):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline(paragraph, line, color=MUTED, size=9.5)
            index += 1
            continue
        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if not candidate.strip():
                index += 1
                break
            if candidate.startswith(("#", "- ", "> ", "```", "|", "[PAGEBREAK]", "[DIAGRAM:")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate.strip())
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    root = Path(__file__).resolve().parent
    repository = root.parents[1]
    source = root / "DOCSYS-UCTAN-UCA-KILAVUZ.md"
    output = repository / "output" / "docx" / "DocSys-Uctan-Uca-Mimari-Isletim-Kullanim-Kilavuzu.docx"
    asset_dir = root / ".generated-assets"
    if len(sys.argv) > 1:
        output = Path(sys.argv[1]).resolve()
    build(source, output, asset_dir)
    print(output)


if __name__ == "__main__":
    main()
